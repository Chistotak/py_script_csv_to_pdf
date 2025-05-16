# release_notes_generator/docx_creator.py
import logging
from datetime import datetime
import os
import re

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
except ImportError:
    logging.critical("Библиотека python-docx не установлена. Установите ее командой: pip install python-docx")
    raise

logger = logging.getLogger(__name__)


def sanitize_text_docx(text):
    """ Базовая очистка текста для DOCX. """
    if text is None:
        return ""
    return str(text)


def get_style_value(style_config, section, key, default_value, value_type=str):
    """Безопасно извлекает и преобразует значение из словаря конфигурации стилей."""
    # Получаем словарь для секции, или пустой словарь, если секции нет
    section_dict = style_config.get(section, {})
    # Получаем строковое значение ключа из этой секции, или дефолтное значение, преобразованное в строку
    val_str = section_dict.get(key, str(default_value))

    try:
        if value_type == int: return int(val_str)
        if value_type == float: return float(val_str)
        if value_type == bool: return val_str.lower() == 'true'
        if value_type == RGBColor:
            hex_color = val_str.lstrip('#')
            if len(hex_color) == 6: return RGBColor.from_string(hex_color)
            logger.warning(
                f"Некорректный HEX цвета '{val_str}' для {section}/{key}. Используется дефолт (черный, если не указан).")
            return default_value if isinstance(default_value, RGBColor) else RGBColor(0, 0, 0)
        return val_str  # str по умолчанию
    except (ValueError, TypeError) as e:  # TypeError на случай, если default_value нельзя преобразовать в str
        logger.warning(
            f"Некорректное значение '{val_str}' или дефолт '{default_value}' для {section}/{key}: {e}. Используется исходный строковый дефолт или пустая строка.")
        return str(default_value) if not isinstance(default_value, (type(None), RGBColor)) else \
            (RGBColor(0, 0, 0) if isinstance(default_value, RGBColor) else "")


def _apply_run_formatting(run, font_name, font_size_pt, bold, italic, color=None, underline=False):
    """Применяет форматирование к объекту Run."""
    if font_name:
        try:
            run.font.name = font_name
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception as e:
            logger.warning(f"Не удалось установить шрифт '{font_name}': {e}")
    if font_size_pt:
        run.font.size = Pt(font_size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    run.underline = underline


def _add_formatted_paragraph(document, text, style_config,
                             font_key='main', fontsize_key='normal_style_base', color_key='normal_text',
                             bold=False, italic=False, underline=False,
                             alignment=None, left_indent_inches=None,
                             space_before_key=None, space_after_key='normal_paragraph_after',
                             line_spacing_rule=None, line_spacing_val=None,
                             keep_with_next=False, keep_together=False, style_name=None):
    """Улучшенная функция добавления параграфа с форматированием из style_config."""
    p = document.add_paragraph()

    # Если text is None, это специальный маркер для параграфа-отступа
    # В этом случае run_text_to_add будет содержать неразрывный пробел
    is_spacing_paragraph = (text is None)
    run_text_to_add = sanitize_text_docx(text) if not is_spacing_paragraph else u'\u00A0'

    font_name_val = get_style_value(style_config, 'Fonts', font_key, 'Arial')
    font_size_pt_val = get_style_value(style_config, 'FontSizes', fontsize_key, 11, value_type=int)

    # Для параграфа-отступа используем минимальный размер шрифта
    actual_font_size_pt = 1 if is_spacing_paragraph else font_size_pt_val

    color_hex_val = get_style_value(style_config, 'Colors', color_key, "000000")
    try:
        run_color_val = RGBColor.from_string(color_hex_val.lstrip('#'))
    except ValueError:
        logger.warning(f"Некорректный HEX '{color_hex_val}' для '{color_key}'. Используется черный.")
        run_color_val = RGBColor(0, 0, 0)

    if style_name:
        try:
            p.style = style_name
            # Форматируем run даже если стиль применен, чтобы переопределить/дополнить
            # Если текст не None (не просто отступ)
            if not is_spacing_paragraph or run_text_to_add.strip():  # Добавляем run, если есть текст
                current_run = p.runs[0] if p.runs and p.runs[0].text else p.add_run()
                current_run.text = run_text_to_add
                _apply_run_formatting(current_run, font_name_val, actual_font_size_pt, bold, italic, run_color_val,
                                      underline)
        except KeyError:
            logger.warning(f"Стиль '{style_name}' не найден. Используется ручное форматирование.")
            # Если стиль не найден, но текст есть
            if not is_spacing_paragraph or run_text_to_add.strip():
                _apply_run_formatting(p.add_run(run_text_to_add), font_name_val, actual_font_size_pt, bold, italic,
                                      run_color_val, underline)
    elif not is_spacing_paragraph or run_text_to_add.strip():  # Стиль не указан, но текст есть (или это не просто отступ)
        _apply_run_formatting(p.add_run(run_text_to_add), font_name_val, actual_font_size_pt, bold, italic,
                              run_color_val, underline)
    elif is_spacing_paragraph:  # Стиль не указан, и это параграф-отступ
        run = p.add_run(run_text_to_add)  # Добавляем неразрывный пробел
        run.font.size = Pt(1)  # Делаем его невидимым

    p_fmt = p.paragraph_format
    if alignment: p_fmt.alignment = alignment
    if left_indent_inches is not None: p_fmt.left_indent = Inches(left_indent_inches)

    # Отступы применяются всегда, даже для "пустых" параграфов
    if space_before_key: p_fmt.space_before = Pt(
        get_style_value(style_config, 'Spacing', space_before_key, 0, value_type=int))
    # Если space_after_key не указан, используем дефолтный 0 для "пустых" параграфов, иначе из конфига
    default_after = 0 if is_spacing_paragraph and not space_after_key else 6
    p_fmt.space_after = Pt(
        get_style_value(style_config, 'Spacing', space_after_key if space_after_key else 'normal_paragraph_after',
                        default_after, value_type=int))

    if line_spacing_rule and line_spacing_val is not None:
        p_fmt.line_spacing_rule = line_spacing_rule
        if line_spacing_rule in [WD_LINE_SPACING.MULTIPLE, WD_LINE_SPACING.AT_LEAST, WD_LINE_SPACING.EXACTLY]:
            p_fmt.line_spacing = float(line_spacing_val)  # line_spacing ожидает float для MULTIPLE
    p_fmt.keep_with_next = keep_with_next
    p_fmt.keep_together = keep_together
    return p


def extract_microservice_info_for_summary_table(grouped_data_keys, main_config_data):
    logger_func = logging.getLogger(__name__)  # Используем логгер этого модуля
    microservices_summary = []
    # Паттерн должен соответствовать ключам в grouped_data_keys, например, "AM2.3.3"
    version_pattern = re.compile(r"^([A-Z]{2})(\d+(\.\d+){1,2})$")
    seen_summary_entries = set()

    for original_ms_key in sorted(list(grouped_data_keys)):
        service_name_for_table = original_ms_key
        version_number_part = ""
        match = version_pattern.match(original_ms_key)
        if match:
            prefix_key_for_config = match.group(1).upper()  # "AM"
            version_number_part = match.group(2)  # "2.3.3"

            # Получаем шаблон имени из main_config_data
            template = main_config_data.get('MicroserviceVersions', {}).get(prefix_key_for_config)
            if template:
                # Извлекаем "чистое" имя сервиса из шаблона
                service_name_for_table = template.replace("{{version}}", "").replace("(версия )", "").strip()
            else:
                # Если шаблона нет, используем префикс как имя сервиса
                service_name_for_table = prefix_key_for_config
        else:
            logger_func.warning(
                f"Не удалось разобрать ключ микросервиса '{original_ms_key}' для сводной таблицы, будет использован как есть.")
            # В этом случае service_name_for_table = original_ms_key, version_number_part = ""

        summary_tuple = (service_name_for_table, version_number_part)
        if summary_tuple not in seen_summary_entries:
            microservices_summary.append({
                'service_name': service_name_for_table,
                'version_number': version_number_part
            })
            seen_summary_entries.add(summary_tuple)
            logger_func.debug(
                f"Для сводной таблицы: Сервис='{service_name_for_table}', Версия='{version_number_part}' (из ключа '{original_ms_key}')")

    return microservices_summary


def create_release_notes_docx(output_filename, title, grouped_data, use_issue_type_grouping,
                              microservices_summary_data=None,
                              main_config=None, style_config=None):
    document = Document()
    logger.info(f"Создание DOCX документа: {output_filename}")

    # Базовые настройки из style_config
    s_font_main = get_style_value(style_config, 'Fonts', 'main', 'Arial')
    s_fontsize_normal_base = get_style_value(style_config, 'FontSizes', 'normal_style_base', 11, value_type=int)
    s_space_after_normal = get_style_value(style_config, 'Spacing', 'normal_paragraph_after', 6, value_type=int)

    try:  # Настройка стиля 'Normal'
        normal_style = document.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = s_font_main
        r_normal = normal_font._element
        r_normal.rPr.rFonts.set(qn('w:eastAsia'), s_font_main);
        r_normal.rPr.rFonts.set(qn('w:cs'), s_font_main)
        normal_font.size = Pt(s_fontsize_normal_base)
        normal_style.paragraph_format.space_after = Pt(s_space_after_normal)
        normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal_style.paragraph_format.line_spacing = 1.15
        logger.info(f"Стиль 'Normal' настроен: шрифт '{s_font_main}' {s_fontsize_normal_base}pt.")
    except Exception as e:
        logger.warning(f"Не удалось настроить стиль 'Normal': {e}.")

    # Логотип (использует main_config для пути и ширины)
    if main_config:
        logo_path_key = main_config.get('General', {}).get('logo_path')
        if logo_path_key:
            actual_logo_path = logo_path_key
            # Определение пути к лого относительно директории основного конфига
            config_dir = main_config.get('_config_dir_', os.getcwd())
            if not os.path.isabs(actual_logo_path):
                actual_logo_path = os.path.join(config_dir, actual_logo_path)

            if os.path.exists(actual_logo_path):
                try:
                    logo_width = get_style_value(main_config, 'General', 'logo_width_inches', 1.5, value_type=float)
                    p_logo = document.add_paragraph();
                    p_logo.add_run().add_picture(actual_logo_path, width=Inches(logo_width))
                    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_logo.paragraph_format.space_after = Pt(12)
                    logger.info(f"Логотип '{actual_logo_path}' добавлен.")
                except Exception as e:
                    logger.error(f"Не удалось добавить логотип '{actual_logo_path}': {e}")
            else:
                logger.warning(f"Файл логотипа '{actual_logo_path}' (из ключа '{logo_path_key}') не найден.")
        else:
            logger.debug("Путь к логотипу (logo_path) не указан в конфигурации.")

    # Заголовок и дата
    _add_formatted_paragraph(document, title, style_config, font_key='title', fontsize_key='title', color_key='title',
                             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after_key='after_title')
    _add_formatted_paragraph(document, f"Дата генерации: {datetime.now().strftime('%Y-%m-%d %H:%M')}", style_config,
                             font_key='main', fontsize_key='date', color_key='date_text', italic=True,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after_key='after_date')

    # Таблица микросервисов
    if microservices_summary_data:
        _add_formatted_paragraph(document, "Состав релиза по микросервисам:", style_config, font_key='section_header',
                                 fontsize_key='summary_table_title', color_key='summary_table_title', bold=True,
                                 space_after_key='after_summary_table_title', keep_with_next=True)
        if microservices_summary_data:  # Список не пустой
            table = document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.columns[0].width = Inches(
                get_style_value(style_config, 'TableLayout', 'summary_table_col1_width_inches', 4.0, value_type=float))
            table.columns[1].width = Inches(
                get_style_value(style_config, 'TableLayout', 'summary_table_col2_width_inches', 1.5, value_type=float))

            hdr_font = get_style_value(style_config, 'Fonts', 'main',
                                       'Arial')  # Используем основной шрифт для заголовков таблицы
            hdr_fontsize = get_style_value(style_config, 'FontSizes', 'summary_table_header', 11, value_type=int)
            hdr_color = get_style_value(style_config, 'Colors', 'table_header_text', RGBColor(0, 0, 0),
                                        value_type=RGBColor)
            hdr_cells = table.rows[0].cells;
            col_texts = ['Микросервис', 'Версия']
            for i, cell_text in enumerate(col_texts):
                p = hdr_cells[i].paragraphs[0];
                p.text = "";
                run = p.add_run(cell_text)
                _apply_run_formatting(run, hdr_font, hdr_fontsize, True, False, color=hdr_color)

            cell_font = get_style_value(style_config, 'Fonts', 'main', 'Arial')
            cell_fontsize = get_style_value(style_config, 'FontSizes', 'summary_table_text', 10, value_type=int)
            cell_color = get_style_value(style_config, 'Colors', 'table_text', RGBColor(0, 0, 0), value_type=RGBColor)
            for item in microservices_summary_data:
                row_cells = table.add_row().cells;
                texts_to_add = [item.get('service_name', 'N/A'), item.get('version_number', 'N/A')]
                for i, cell_content in enumerate(texts_to_add):
                    row_cells[i].text = cell_content  # Сначала задаем текст ячейки
                    # Форматируем первый параграф и первый run в ячейке
                    if row_cells[i].paragraphs:
                        p_cell = row_cells[i].paragraphs[0]
                        # Убедимся, что run существует и содержит текст
                        current_run = p_cell.runs[0] if p_cell.runs else p_cell.add_run(cell_content)
                        if not current_run.text and cell_content: current_run.text = cell_content  # Если run был, но пустой
                        _apply_run_formatting(current_run, cell_font, cell_fontsize, False, False, color=cell_color)
            logger.info("Таблица микросервисов добавлена.")
        else:
            _add_formatted_paragraph(document, "Информация о версиях микросервисов в данном релизе отсутствует.",
                                     style_config, fontsize_key='normal_style_base', italic=True)
        _add_formatted_paragraph(document, None, style_config, space_after_key='after_summary_table')
    elif grouped_data:
        _add_formatted_paragraph(document, None, style_config, space_after_key='normal_paragraph_after')

    # Детализация задач
    if not grouped_data:
        if not microservices_summary_data:
            _add_formatted_paragraph(document, "Нет задач для отображения.", style_config, font_key='main',
                                     fontsize_key='normal_style_base')
    else:
        microservice_version_pattern_docx = re.compile(r"^([A-Z]{2})(\d+(\.\d+){1,2})$")
        sorted_ms_versions_original = sorted(grouped_data.keys())

        for ms_idx, ms_version_original_key in enumerate(sorted_ms_versions_original):
            display_ms_version = ms_version_original_key
            if main_config:
                match = microservice_version_pattern_docx.match(ms_version_original_key)
                if match:
                    prefix_to_lookup = match.group(1).upper()
                    version_number = match.group(2)
                    template = main_config.get('MicroserviceVersions', {}).get(prefix_to_lookup)
                    if template: display_ms_version = template.replace("{{version}}", version_number)

            _add_formatted_paragraph(document, display_ms_version, style_config,
                                     font_key='section_header', fontsize_key='ms_version_header',
                                     color_key='section_header', bold=True,
                                     space_before_key='section_after_space' if ms_idx > 0 else None,
                                     space_after_key='after_ms_version_header', keep_with_next=True)

            data_for_ms = grouped_data[ms_version_original_key]
            if use_issue_type_grouping and isinstance(data_for_ms, dict):
                sorted_issue_types_display = sorted(data_for_ms.keys())
                for type_idx, issue_type_display_name in enumerate(sorted_issue_types_display):
                    _add_formatted_paragraph(document, issue_type_display_name, style_config,
                                             font_key='issue_type_header', fontsize_key='issue_type_header',
                                             color_key='sub_header', bold=True,
                                             left_indent_inches=0.25,
                                             space_before_key='after_ms_version_header',
                                             space_after_key='after_issue_type_header', keep_with_next=True)

                    tasks_to_process = data_for_ms[issue_type_display_name]
                    current_base_indent = 0.50
                    for task_num, task in enumerate(tasks_to_process):
                        _add_formatted_paragraph(document, task['key'] + ":", style_config,
                                                 font_key='task_key', fontsize_key='task_key',
                                                 color_key='task_key', bold=True,
                                                 left_indent_inches=current_base_indent,
                                                 space_before_key='task_block_internal_space',
                                                 space_after_key='task_key_after')

                        cust_desc_text = sanitize_text_docx(task['cust_desc'])
                        is_cust_desc_empty = not bool(task['cust_desc'])
                        _add_formatted_paragraph(document,
                                                 cust_desc_text if not is_cust_desc_empty else "Описание для клиента отсутствует.",
                                                 style_config, font_key='main', fontsize_key='task_description',
                                                 color_key='task_description', italic=is_cust_desc_empty,
                                                 left_indent_inches=current_base_indent,
                                                 space_after_key='task_description_after' if not task[
                                                     'install_instr'] else 'task_block_internal_space')

                        if task['install_instr']:
                            _add_formatted_paragraph(document, "Инструкция по установке:", style_config,
                                                     font_key='main', fontsize_key='install_instruction_label',
                                                     color_key='install_instruction_label', bold=True,
                                                     left_indent_inches=current_base_indent,
                                                     space_before_key='task_block_internal_space',
                                                     space_after_key='install_label_after')
                            _add_formatted_paragraph(document, sanitize_text_docx(task['install_instr']), style_config,
                                                     font_key='main', fontsize_key='install_instruction_text',
                                                     color_key='install_instruction_text',
                                                     left_indent_inches=current_base_indent,
                                                     space_after_key='install_text_after')
            elif isinstance(data_for_ms, list):  # Задачи без группировки по типу
                current_base_indent = 0.25
                tasks_to_process = data_for_ms
                for task_num, task in enumerate(tasks_to_process):
                    _add_formatted_paragraph(document, task['key'] + ":", style_config, font_key='task_key',
                                             fontsize_key='task_key', color_key='task_key', bold=True,
                                             left_indent_inches=current_base_indent,
                                             space_before_key='task_block_internal_space',
                                             space_after_key='task_key_after')
                    cust_desc_text = sanitize_text_docx(task['cust_desc'])
                    is_cust_desc_empty = not bool(task['cust_desc'])
                    _add_formatted_paragraph(document,
                                             cust_desc_text if not is_cust_desc_empty else "Описание для клиента отсутствует.",
                                             style_config, font_key='main', fontsize_key='task_description',
                                             color_key='task_description', italic=is_cust_desc_empty,
                                             left_indent_inches=current_base_indent,
                                             space_after_key='task_description_after' if not task[
                                                 'install_instr'] else 'task_block_internal_space')
                    if task['install_instr']:
                        _add_formatted_paragraph(document, "Инструкция по установке:", style_config, font_key='main',
                                                 fontsize_key='install_instruction_label',
                                                 color_key='install_instruction_label', bold=True,
                                                 left_indent_inches=current_base_indent,
                                                 space_before_key='task_block_internal_space',
                                                 space_after_key='install_label_after')
                        _add_formatted_paragraph(document, sanitize_text_docx(task['install_instr']), style_config,
                                                 font_key='main', fontsize_key='install_instruction_text',
                                                 color_key='install_instruction_text',
                                                 left_indent_inches=current_base_indent,
                                                 space_after_key='install_text_after')

            if ms_idx < len(sorted_ms_versions_original) - 1:  # Отступ между секциями микросервисов
                _add_formatted_paragraph(document, None, style_config, space_after_key='section_after_space')

    try:
        document.save(output_filename)
        logger.info(f"DOCX '{output_filename}' успешно сохранен.")
        return True
    except Exception as e:
        logger.error(f"Ошибка при сохранении DOCX '{output_filename}': {e}", exc_info=True)
        return False